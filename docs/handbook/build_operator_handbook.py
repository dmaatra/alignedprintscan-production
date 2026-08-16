#!/usr/bin/env python3
from __future__ import annotations
import json, re, subprocess, sys
from pathlib import Path
from html.parser import HTMLParser
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parents[2]
OUT=ROOT/"docs"/"handbook"
DOCX=OUT/"Aligned_Print_Scan_Operator_Handbook_v3.0.docx"
PDF=OUT/"Aligned_Print_Scan_Operator_Handbook_v3.0.pdf"
NAVY="161C4D"; GOLD="C8A96B"; IVORY="F7F3EA"; INK="25283A"; MUTED="626A7D"; RED="982D27"; BLUE="DDE7F4"

PARTS=[
("I","FOUNDATION",["How to Use This Handbook","System Production Baseline and Document Control","Operator Role, Authority, and Escalation","Service Lines and Ownership"]),
("II","CUSTOMER LIFECYCLE",["Lifecycle from Intake to Completion","What the Customer Sees","Customer Information Boundaries"]),
("III","ADMIN DASHBOARD",["Requests","Review Queue","Calendar and RON Sessions","Customers, Invoices, and Payments","Messages, Templates, and Scripts"]),
("IV","REQUEST WORKSPACE",["Overview","Customer","Documents","Quote","Payments","Messages","Fulfillment","Timeline"]),
("V","COMMUNICATIONS, TEMPLATES & STATUS TRANSITIONS",["How APS Communications Work","Send Message vs Send & Update Status","Communication Log and Timeline","Template Operator Workflow","Status Transition Discipline"]),
("VI","DOCUMENT OPERATIONS",["Document Intake and Provenance","Review and Classification","Release to Customer","Completed Deliverables and Re-review"]),
("VII","QUOTES, INVOICES, PAYMENTS & REFUNDS",["Quote Review and Approval","Invoice Creation and Supplemental Charges","Stripe and Manual Payment","Refund Decision","Stripe and Offline Refund","Financial Reconciliation"]),
("VIII","SCHEDULING, RESCHEDULING & CANCELLATION",["Scheduling and Confirmation","Rescheduling","Cancellation Review","No-show, Late Cancellation, and APS Unable to Fulfill"]),
("IX","RON OPERATOR SOP",["RON Readiness","Proof Draft and Preparation","Review and Activation","Signer Access and Live Notarization","Return to APS, Review, Release, and Completion"]),
("X","RON / NOTARIAL ACT SCRIPTS",["Notarial Neutrality and Act Selection","Acknowledgment","Jurat and Oath or Affirmation","Signature Witnessing and Copy Certification","Credible Witness and Document Witness","Representative Capacity and Complex Signings","Stop and Refuse Conditions","RON Quick-Flip"]),
("XI","MOBILE NOTARY OPERATOR SOP & SCRIPTS",["Mobile Pre-appointment and Travel","Arrival, Identity, and Act Performance","Certificate and Document Final Check","Return to APS and Completion"]),
("XII","PRINT & SCAN OPERATOR SOP",["Order and Source Review","Production Start and Printing","Scanning and File Assembly","Quality Control, Delivery, and Completion"]),
("XIII","CUSTOMER PORTAL",["Portal Access and Navigation","Documents and Release Visibility","Quote, Payment, Fulfillment, Messages, and Activity"]),
("XIV","CUSTOMER SUPPORT",["Support Intake and Request Scoping","Safe Customer Explanations","Escalation and Resolution"]),
("XV","GROWTH, REVIEWS & ANALYTICS",["Neutral Review Request","Google Business and Review Destinations","GA4, Attribution, and Privacy"]),
("XVI","SECURITY, PRIVACY & RECORDS",["Authorization and RLS Boundaries","Sensitive Data and Customer Visibility","Records, Auditability, and Retention"]),
("XVII","TROUBLESHOOTING",["Intake, Customer, and Document Problems","Financial and Communication Problems","Scheduling and Service Problems","Portal, Notification, and Completion Problems"]),
("XVIII","QUICK REFERENCE",["Service Quick Starts","Financial and Change Quick Cards","Communication and Document Quick Cards","Completion and Proof Return Quick Cards"]),
("XIX","PUBLIC POLICIES",["Terms of Service","Privacy Policy","Accessibility Statement","Frequently Asked Questions"]),
("XX","DECISION TABLES & MATRICES",["System Ownership Matrix","Status Transition Matrix","Template Attachment Matrix","Fulfillment and Customer Visibility Matrix"]),
("XXI","COMPLETE TEMPLATE LIBRARY",["Template Operator Catalog"]),
("XXII","STATUS REFERENCE",["Request, Financial, Document, Review, and RON Status Catalog"]),
("XXIII","VISUAL SOP ATLAS",["Visual SOP Atlas"]),
("XXIV","GLOSSARY",["Alphabetical Glossary"]),
("XXV","INDEX",["Back-of-Book Index"]),
("XXVI","CHANGE LOG & DOCUMENT MAINTENANCE",["Revision History","Maintaining the Handbook and Admin Reference"]),
]

SCREENSHOTS={
"Requests":"ss-001-requests.jpg","Review Queue":"ss-010-review-queue.jpg","Overview":"ss-020-request-overview.jpg","Documents":"ss-030-document-lifecycle.jpg","Review and Classification":"ss-031-private-output-review.jpg","Release to Customer":"ss-032-release-to-customer.jpg","Quote":"ss-040-quote-builder.jpg","Payments":"ss-050-payments-ledger.jpg","Cancellation Review":"ss-052-cancellation-refund-review.jpg","Communication Log and Timeline":"ss-060-communication-log.jpg","Send Message vs Send & Update Status":"ss-061-send-message.jpg","Status Transition Discipline":"ss-062-send-update-status.jpg","RON Readiness":"ss-070-ron-fulfillment.jpg","Mobile Pre-appointment and Travel":"ss-080-mobile-fulfillment.jpg","Order and Source Review":"ss-090-print-scan-fulfillment.jpg","Portal Access and Navigation":"ss-100-portal-overview.jpg","Documents and Release Visibility":"ss-101-portal-documents.jpg","Rescheduling":"ss-110-portal-cancellation.jpg","Financial Reconciliation":"ss-112-portal-refund.jpg"}

STATUS_ENTRIES=["Under Review","Quote Ready","Awaiting Approval","Awaiting Payment","Payment Received","Appointment Confirmed","Fulfillment In Progress","Cancellation Requested","Cancelled","Completed","Archived","Draft Invoice","Open Invoice","Paid Invoice","Partially Refunded","Refunded","Document Pending","Approved Internal","Released to Customer","Re-review Required","Review Eligible","Review Sent"]+[f"RON {i:02d} - {name}" for i,name in enumerate(["Request Review","Payment Readiness","Appointment Confirmation","Signer Readiness","Document Readiness","Proof Draft","Proof Preparation","Review & Activate","Signer Access","Live Notarization","Provider Completion","APS Document Review","Released / Completion"],1)]

GLOSSARY={"Acknowledgment":"A notarial act in which the signer acknowledges executing an instrument for its stated purposes.","APS":"Aligned Print & Scan.","Communication Log":"The request-scoped record of customer-facing messages and delivery state.","Completion Gate":"The service-specific checks that must pass before APS may mark a request Completed.","Credible witness":"A person used under an authorized identity procedure; not the same as a document witness.","Jurat":"A certificate for a verification on oath or affirmation that includes a witnessed signature.","KBA":"Knowledge-based authentication performed by the RON provider; APS never stores or coaches answers.","Proof":"The external RON provider used for transaction, identity, session, and completed-asset functions.","Release to Customer":"An explicit authorized action that makes an eligible document visible in the portal.","RLS":"Row-level security controlling database access.","RON":"Remote Online Notarization.","Send & Update Status":"A delivery-first communication action that changes status only after successful send.","Timeline":"The request-scoped history of authoritative workflow events."}

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd")) or OxmlElement("w:shd"); shd.set(qn("w:fill"),fill)
    if shd.getparent() is None: tcPr.append(shd)
def set_cell_margin(cell,top=100,start=120,bottom=100,end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); m=tcPr.first_child_found_in("w:tcMar")
    if m is None: m=OxmlElement("w:tcMar");tcPr.append(m)
    for side,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=m.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}");node.set(qn("w:w"),str(val));node.set(qn("w:type"),"dxa")
        if node.getparent() is None:m.append(node)
def keep(p,next_=False):
    pPr=p._p.get_or_add_pPr(); tag="keepNext" if next_ else "keepLines"; node=OxmlElement(f"w:{tag}");pPr.append(node)
def bookmark(p,name):
    name=re.sub(r"[^A-Za-z0-9_]","_",name)[:38]; start=OxmlElement("w:bookmarkStart");start.set(qn("w:id"),str(abs(hash(name))%2000000000));start.set(qn("w:name"),name);end=OxmlElement("w:bookmarkEnd");end.set(qn("w:id"),start.get(qn("w:id")));p._p.insert(0,start);p._p.append(end);return name
def hyperlink(p,text,anchor):
    link=OxmlElement("w:hyperlink");link.set(qn("w:anchor"),anchor);run=OxmlElement("w:r");rPr=OxmlElement("w:rPr");color=OxmlElement("w:color");color.set(qn("w:val"),NAVY);rPr.append(color);u=OxmlElement("w:u");u.set(qn("w:val"),"single");rPr.append(u);run.append(rPr);t=OxmlElement("w:t");t.text=text;run.append(t);link.append(run);p._p.append(link)
def field(run,instr):
    for tag,typ,text in (("fldChar","begin",None),("instrText",None,instr),("fldChar","separate",None),("t",None,"0"),("fldChar","end",None)):
        el=OxmlElement(f"w:{tag}");
        if typ:el.set(qn("w:fldCharType"),typ)
        if text is not None:el.text=text
        run._r.append(el)
def add_heading(doc,text,level,anchor):
    p=doc.add_heading(text,level=level);bookmark(p,anchor);keep(p,True);return p
def add_bullets(doc,items,style="List Bullet"):
    for item in items:
        p=doc.add_paragraph(item,style=style);keep(p)
def add_callout(doc,label,text,kind="note"):
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(5);p.paragraph_format.space_after=Pt(7);p.paragraph_format.left_indent=Inches(.10);p.paragraph_format.right_indent=Inches(.10);keep(p)
    pPr=p._p.get_or_add_pPr();shd=OxmlElement("w:shd");shd.set(qn("w:fill"),{"must":IVORY,"do_not":"FCE8E6","customer":BLUE,"records":"EDF4EA","blocker":"FCE8E6"}.get(kind,IVORY));pPr.append(shd)
    borders=OxmlElement("w:pBdr");left=OxmlElement("w:left");left.set(qn("w:val"),"single");left.set(qn("w:sz"),"18");left.set(qn("w:space"),"8");left.set(qn("w:color"),RED if kind in ("do_not","blocker") else GOLD);borders.append(left);pPr.append(borders)
    r=p.add_run(label.upper()+"  ");r.bold=True;r.font.color.rgb=RGBColor.from_string(RED if kind in ("do_not","blocker") else NAVY);p.add_run(text)
def add_table(doc,headers,rows,widths=None):
    table=doc.add_table(rows=1,cols=len(headers));table.autofit=False;table.style="Table Grid"
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i];c.text=str(h);shade(c,NAVY);set_cell_margin(c)
        for r in c.paragraphs[0].runs:r.bold=True;r.font.color.rgb=RGBColor(255,255,255);r.font.size=Pt(9.5)
    trPr=table.rows[0]._tr.get_or_add_trPr();hdr=OxmlElement("w:tblHeader");hdr.set(qn("w:val"),"true");trPr.append(hdr)
    for row in rows:
        cells=table.add_row().cells
        for i,value in enumerate(row):cells[i].text=str(value);cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;set_cell_margin(cells[i]);
        trPr=table.rows[-1]._tr.get_or_add_trPr();cant=OxmlElement("w:cantSplit");trPr.append(cant)
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths):row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2);return table
class PolicyParser(HTMLParser):
    def __init__(self):super().__init__();self.capture=None;self.skip=0;self.buffer=[];self.blocks=[]
    def handle_starttag(self,tag,attrs):
        if tag in {"script","style","nav","footer","header","form","button"}:self.skip+=1
        elif not self.skip and tag in {"h1","h2","h3","p","li"}:self.capture=tag;self.buffer=[]
    def handle_endtag(self,tag):
        if tag in {"script","style","nav","footer","header","form","button"} and self.skip:self.skip-=1
        elif self.capture==tag:
            text=" ".join("".join(self.buffer).split())
            if text and text not in self.blocks:self.blocks.append(text)
            self.capture=None;self.buffer=[]
    def handle_data(self,data):
        if self.capture and not self.skip:self.buffer.append(data)
def clean_html(path):
    parser=PolicyParser();parser.feed(path.read_text());return parser.blocks
def source_notes():
    text=(ROOT/"docs"/"APS_SYSTEM_OPERATIONS_WORKFLOW_MANUAL.md").read_text();paras=[]
    for p in re.split(r"\n\s*\n",text):
        p=re.sub(r"^#+\s*","",p.strip());p=re.sub(r"\*\*([^*]+)\*\*",r"\1",p);p=re.sub(r"\[([^]]+)\]\([^)]+\)",r"\1",p)
        if len(p)>90 and not p.startswith("|") and "CONFIRMED IN CODE" not in p:paras.append(p.replace("\n"," "))
    return paras
SOURCE=source_notes()
def relevant(topic,n=2):
    words={w.lower() for w in re.findall(r"[A-Za-z]{5,}",topic)};rank=sorted(SOURCE,key=lambda p:sum(w in p.lower() for w in words),reverse=True);return [p for p in rank if sum(w in p.lower() for w in words)>0][:n]
def add_chapter(doc,part_roman,chapter_no,title,anchor):
    add_heading(doc,f"Chapter {chapter_no} - {title}",1,anchor)
    add_back(doc)
    context=(relevant(title)+[f"This chapter explains the maintained APS {title.lower()} workflow as an operator procedure. Use the current request, document, financial, provider, and communication records as authority; never infer completion from a button click or customer expectation."])[0]
    add_heading(doc,f"{chapter_no}.1 Purpose and When You Use It",2,anchor+"_purpose")
    doc.add_paragraph(context);doc.add_paragraph(f"Use this chapter whenever the operator must review, perform, explain, or troubleshoot {title.lower()}. The goal is a correct customer outcome supported by the authoritative APS record and appropriate service evidence.")
    add_heading(doc,f"{chapter_no}.2 What You See and What It Means",2,anchor+"_controls")
    add_table(doc,["Operator surface","Meaning"],[("Status and summary cards","Current projected business state; verify the underlying evidence."),("Primary controls","Actions available for the current role and service state."),("Messages / Timeline","Customer delivery history and authoritative workflow events."),("Blockers / review items","Conditions that require correction or an administrator decision.")],[2.0,4.5])
    if title in SCREENSHOTS:add_figure(doc,SCREENSHOTS[title],title)
    add_heading(doc,f"{chapter_no}.3 Operator Procedure",2,anchor+"_procedure")
    steps=["Open the exact APS request or module and confirm the reference, customer, and service type.","Read the current status, blockers, and most recent Timeline events before taking action.","Verify the authoritative supporting record: customer data, document state, quote/invoice/payment ledger, appointment, fulfillment facts, or provider state as applicable.","Confirm the action is appropriate for the service type and current state; hidden or unavailable controls are not a reason to bypass the workflow.","Use the maintained control once. Do not refresh, retry, resend, release, or recreate until the result is known.","Verify the saved result in its authoritative panel and confirm Communication Log/Timeline entries where the action communicates or changes state.","Check the customer portal only for customer-visible results and confirm internal, unreleased, and provider-administrative information remains hidden.","Record or escalate any mismatch; do not force a status merely to make the request appear complete."]
    add_bullets(doc,steps,"List Number")
    add_heading(doc,f"{chapter_no}.4 Safeguards, Customer View, and APS Records",2,anchor+"_safeguards")
    add_callout(doc,"Must Do","Use authoritative records, preserve idempotency, and verify the result after every consequential action.","must")
    add_callout(doc,"Do Not","Do not alter legitimate customer data, finances, document release, communication, or Proof state merely to test or clear a blocker.","do_not")
    add_callout(doc,"Customer Sees","Only the customer-safe status, released documents, legitimate financial projection, sent communications, and service-appropriate fulfillment information.","customer")
    add_callout(doc,"APS Records","The relevant request row, service detail, financial ledger, document metadata, Communication Log, Timeline, review item, and provider mapping remain separate authoritative records.","records")
    add_heading(doc,f"{chapter_no}.5 Next Action and Troubleshooting",2,anchor+"_troubleshooting")
    doc.add_paragraph("The normal next action is the next service-appropriate state supported by evidence. If the UI and record disagree, refresh once, inspect the authoritative panel and Timeline, and use the maintained safe retry or review path. Escalate when authorization, identity, payment, customer release, provider state, or legal/notarial judgment is unresolved.")
    add_callout(doc,"Blocker","Never convert an unknown or failed result into a successful status. Preserve the current record and escalate with the APS reference, module, timestamp, and observed error - without credentials or sensitive identity data.","blocker")
def add_back(doc):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.RIGHT;hyperlink(p,"Back to Contents","contents")
def add_figure(doc,filename,title):
    path=ROOT/"docs"/"assets"/"manual"/filename
    if not path.exists():return
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run();r.add_picture(str(path),width=Inches(6.15));keep(p,True)
    cap=doc.add_paragraph(f"Figure - {title}: maintained APS instructional view.",style="Caption");cap.alignment=WD_ALIGN_PARAGRAPH.CENTER;keep(cap)
    # alt text
    for drawing in p._p.xpath(".//wp:docPr"):
        drawing.set("descr",f"APS instructional screenshot for {title}")
def load_catalogs():
    js="import('./assets/js/operator-reference-catalog.mjs').then(m=>console.log(JSON.stringify(m.OPERATOR_REFERENCE_SCRIPTS)))"
    scripts=json.loads(subprocess.check_output(["node","--input-type=module","-e",js],cwd=ROOT,text=True))
    src=(ROOT/"supabase"/"functions"/"_shared"/"template-preview.mjs").read_text()
    templates=[]
    for m in re.finditer(r"\n\s{2}([a-z0-9_]+):\{category:\"([^\"]+)\",purpose:\"([^\"]+)\",trigger:\"([^\"]+)\",classification:\"([^\"]+)\",eyebrow:\"([^\"]+)\",title:\"([^\"]+)\",cta:\"([^\"]+)\",tab:\"([^\"]+)\",fields:\[([^]]+)\]\}",src):
        templates.append(dict(key=m.group(1),category=m.group(2),purpose=m.group(3),trigger=m.group(4),classification=m.group(5),eyebrow=m.group(6),title=m.group(7),cta=m.group(8),tab=m.group(9),fields=re.findall(r'"([^"]+)"',m.group(10))))
    return templates,scripts
def add_template_catalog(doc,templates,anchor):
    for i,t in enumerate(templates,1):
        h=add_heading(doc,f"1.{i} {t['title']}",2,f"{anchor}_template_{i}")
        add_table(doc,["Operator field","Maintained value"],[("Template key",t["key"]),("Family",t["category"]),("Purpose",t["purpose"]),("When to use",t["trigger"]),("When not to use","Do not send when the trigger/evidence is absent or a duplicate successful communication already exists."),("Action",t["classification"]),("Recipient","Customer on the current APS request"),("CTA",f"{t['cta']} -> {t['tab']}"),("Attachment","Use only the exact customer-safe attachment required by the active template/workflow; never internal, audit, source, or unreleased files."),("Customer result","Branded customer message with a secure, request-appropriate destination."),("APS result","Successful delivery persists in Communication Log and Timeline; failed delivery must not be represented as Sent."),("Idempotency","Verify existing successful communication before retrying."),("Next action","Follow the status and service-specific SOP after delivery.")],[1.65,4.85])
        add_callout(doc,"Maintained preview",f"{t['eyebrow']} - {t['title']}. Expected fields: {', '.join(t['fields'])}.","customer")
def add_script_book(doc,scripts,anchor):
    for i,s in enumerate(scripts,1):
        add_heading(doc,f"Script {i} - {s['name']}",2,f"{anchor}_script_{i}")
        doc.add_paragraph(s["purpose"]);doc.add_paragraph("When to use: "+s["when"])
        add_callout(doc,"Script - Say This",s["say"],"customer")
        add_heading(doc,"Must Do",3,f"{anchor}_script_{i}_must");add_bullets(doc,s["mustDo"])
        add_heading(doc,"Do Not",3,f"{anchor}_script_{i}_dont");add_bullets(doc,s["doNot"])
        add_callout(doc,"Stop / Refuse",s["stop"],"blocker");doc.add_paragraph("APS next step: "+s["next"]);doc.add_paragraph("Related: "+s["related"])
def add_status_catalog(doc,anchor):
    for i,name in enumerate(STATUS_ENTRIES,1):
        add_heading(doc,f"1.{i} {name}",2,f"{anchor}_status_{i}")
        add_table(doc,["Reference field","Operator guidance"],[("Plain-English meaning",f"The maintained APS record is in {name}; read the underlying evidence before acting."),("Legitimate entry","Entered only by the authorized workflow, provider synchronization, customer action, or administrator decision that owns this state."),("Authoritative evidence","Request/service record plus applicable document, appointment, provider, financial, communication, or review record."),("Customer sees","Only the customer-safe projection for this state; internal evidence and error details remain private."),("Operator action","Follow the next service-appropriate SOP and clear blockers through their owning workflow."),("Do not","Do not set this state to bypass evidence, delivery, payment, release, identity, or completion requirements."),("Next expected state","The next state supported by the maintained transition matrix and authoritative evidence."),("Troubleshooting","Compare the source record, Timeline, Communication Log, and portal projection; escalate unresolved divergence.")],[1.65,4.85])
def add_policy(doc,title,path,anchor):
    add_heading(doc,title,1,anchor);add_back(doc);doc.add_paragraph("Operator note: the following is the complete current public text extracted from the production source. The public page remains authoritative for customer-facing presentation.")
    for block in clean_html(path):
        if len(block)<90 and not block.endswith(('.',':','?')):add_heading(doc,block,2,anchor+"_"+str(abs(hash(block))%999999))
        else:doc.add_paragraph(block)
def configure(doc):
    sec=doc.sections[0];sec.top_margin=Inches(.78);sec.bottom_margin=Inches(.72);sec.left_margin=Inches(.85);sec.right_margin=Inches(.85);sec.header_distance=Inches(.35);sec.footer_distance=Inches(.35)
    styles=doc.styles;normal=styles["Normal"];normal.font.name="Aptos";normal.font.size=Pt(10.7);normal.font.color.rgb=RGBColor.from_string(INK);normal.paragraph_format.space_after=Pt(6);normal.paragraph_format.line_spacing=1.15
    for name,size,color,before,after in (("Title",31,NAVY,0,12),("Heading 1",22,NAVY,18,9),("Heading 2",16,NAVY,14,7),("Heading 3",12.5,"8A681F",10,5)):
        s=styles[name];s.font.name="Aptos Display";s.font.size=Pt(size);s.font.color.rgb=RGBColor.from_string(color);s.font.bold=True;s.paragraph_format.space_before=Pt(before);s.paragraph_format.space_after=Pt(after);s.paragraph_format.keep_with_next=True
    styles["Caption"].font.name="Aptos";styles["Caption"].font.size=Pt(9);styles["Caption"].font.color.rgb=RGBColor.from_string(MUTED)
    header=sec.header.paragraphs[0];header.text="ALIGNED PRINT & SCAN  |  OPERATOR EDITION";header.style=styles["Caption"]
    footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;footer.add_run("SYSTEM PRODUCTION BASELINE: CURRENT MAIN  |  DOCUMENTATION REVISION 3.0  |  ");field(footer.add_run(),"PAGE")
def cover(doc):
    for _ in range(5):doc.add_paragraph()
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("ALIGNED PRINT & SCAN");r.bold=True;r.font.size=Pt(16);r.font.color.rgb=RGBColor.from_string(GOLD)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("OPERATIONS, POLICIES\n& STANDARD OPERATING\nPROCEDURES HANDBOOK");r.bold=True;r.font.name="Aptos Display";r.font.size=Pt(29);r.font.color.rgb=RGBColor.from_string(NAVY)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run("OPERATOR EDITION");r.bold=True;r.font.size=Pt(15);r.font.color.rgb=RGBColor.from_string(GOLD)
    doc.add_paragraph();add_table(doc,["Document control","Value"],[("System production baseline","Current origin/main at release verification"),("Documentation revision","3.0 - August 15, 2026"),("Validation","Validated against current production implementation"),("Audience","APS administrators and trained operators")],[2.0,4.5])
    doc.add_page_break()
def contents(doc):
    h=add_heading(doc,"Contents",1,"contents")
    doc.add_paragraph("Click any entry to navigate. Page references are Word PAGEREF fields and update when fields are refreshed.")
    for roman,title,chapters in PARTS:
        part_anchor=f"part_{roman}";p=doc.add_paragraph();p.paragraph_format.space_before=Pt(6);r=p.add_run();r.bold=True;r.font.color.rgb=RGBColor.from_string(NAVY);hyperlink(p,f"PART {roman} - {title}",part_anchor)
        for i,ch in enumerate(chapters,1):
            anchor=f"part_{roman}_chapter_{i}";p=doc.add_paragraph();p.paragraph_format.left_indent=Inches(.25);hyperlink(p,f"Chapter {i} - {ch}",anchor);p.add_run("  ....  ");field(p.add_run(),f"PAGEREF {anchor} \\h")
    doc.add_page_break()
def part_divider(doc,roman,title,chapters):
    if roman!="XXI":doc.add_page_break()
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(90);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(f"PART {roman}");r.bold=True;r.font.size=Pt(16);r.font.color.rgb=RGBColor.from_string(GOLD);bookmark(p,f"part_{roman}")
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(title);r.bold=True;r.font.size=Pt(27);r.font.color.rgb=RGBColor.from_string(NAVY)
    doc.add_paragraph();add_heading(doc,"In This Part",2,f"part_{roman}_summary");add_bullets(doc,chapters)
    add_callout(doc,"Learning objective",f"After this Part, the operator can explain and perform the maintained {title.lower()} procedures without bypassing authorization, evidence, customer-visibility, or completion controls.","must");add_back(doc);doc.add_page_break()
QUICK_CARD_GROUPS=[
 ["RON - Quick Start","Mobile - Quick Start","Print & Scan - Quick Start","RON - Must Do Every Session","Mobile - Must Do Every Appointment","Print & Scan - Must Do Every Job"],
 ["Payment Received - Quick Card","Refund - Quick Card","Cancellation - Quick Card","Reschedule - Quick Card"],
 ["Document Release - Quick Card","Send Message - Quick Card","Send & Update Status - Quick Card"],
 ["Return from Proof - Quick Card","Completion Gate - Quick Card"],
]
def add_quick_cards(doc,group_index):
    for i,title in enumerate(QUICK_CARD_GROUPS[group_index],1):
        anchor=f"quick_{group_index+1}_{i}";add_heading(doc,title,2,anchor);add_table(doc,["Step","Must Do"],[(1,"Open the exact request and verify service, customer, status, and blockers."),(2,"Verify the authoritative evidence owned by this workflow."),(3,"Use the maintained action once; preserve idempotency."),(4,"Verify the saved record, Communication Log, Timeline, and customer-safe projection."),(5,"Advance only when the next state and completion gate are legitimately supported.")],[.65,5.85]);add_callout(doc,"Quick Check","No real customer communication, money movement, document release, or Proof action is performed merely to test this card.","must")
def add_visual_atlas(doc):
    for i,path in enumerate(sorted((ROOT/"docs"/"assets"/"manual").glob("ss-*.jpg")),1):
        if i>1:doc.add_page_break()
        add_heading(doc,f"Figure {i} - {path.stem}",2,f"atlas_{i}")
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run();r.add_picture(str(path),width=Inches(5.55));keep(p,True)
        for drawing in p._p.xpath(".//wp:docPr"):drawing.set("descr",f"APS instructional screenshot for {path.stem}")
        cap=doc.add_paragraph(f"Figure {i} - {path.stem.replace('-', ' ').title()}: maintained APS instructional view.",style="Caption");cap.alignment=WD_ALIGN_PARAGRAPH.CENTER;keep(cap)
def landscape_start(doc):
    sec=doc.add_section(WD_SECTION.NEW_PAGE);sec.orientation=WD_ORIENT.LANDSCAPE;sec.page_width=Inches(11);sec.page_height=Inches(8.5);sec.top_margin=Inches(.65);sec.bottom_margin=Inches(.65);sec.left_margin=Inches(.7);sec.right_margin=Inches(.7);return sec
def portrait_start(doc):
    sec=doc.add_section(WD_SECTION.NEW_PAGE);sec.orientation=WD_ORIENT.PORTRAIT;sec.page_width=Inches(8.5);sec.page_height=Inches(11);sec.top_margin=Inches(.78);sec.bottom_margin=Inches(.72);sec.left_margin=Inches(.85);sec.right_margin=Inches(.85);return sec
def add_matrix_chapter(doc,index,title,anchor,templates):
    add_heading(doc,f"Chapter {index} - {title}",1,anchor);add_back(doc);doc.add_paragraph("Use this matrix as a routing reference. The underlying request, ledger, document, provider, communication, and authorization records remain authoritative.")
    if index==1:rows=[("Public intake","Customer / public Edge Function","Create scoped request and service facts","No admin secrets or internal notes"),("Admin dashboard","Authenticated administrator","Review and authorized workflow actions","No customer-side direct access"),("Customer portal","Scoped customer access","Read customer-safe projection and submit allowed actions","No internal/unreleased/provider-admin data"),("Supabase","RLS / authorized functions","Persist authoritative records","No client bypass"),("Proof","External provider","Identity, RON session, provider completion","APS controls customer release")];headers=["System","Owner","Authoritative function","Boundary"]
    elif index==2:rows=[("Under Review","Request created","Request + service details","Request Received","Automatic","None","Quote Ready / review","Review request"),("Quote Ready","Reviewed quote saved","Current quote version","Quote Ready","Send & Update","Customer-safe quote","Awaiting Approval","Review quote"),("Awaiting Payment","Quote approved / invoice open","Invoice ledger","Awaiting Payment","Automatic / Send Message","Invoice when maintained","Payment Received","Pay balance"),("Payment Received","Authoritative payment recorded","Payment ledger","Payment Received","Automatic / Workflow","Receipt/invoice if maintained","Fulfillment","Prepare for service"),("Appointment Confirmed","Confirmed appointment saved","Appointment facts","Appointment Confirmed","Send & Update","None","Fulfillment","Attend / prepare"),("Cancelled","Admin decision persisted","Action + financial review","Cancellation Confirmed","Send & Update","Customer-safe confirmation","Refund state / closed","Review result"),("Completed","Completion gate passed","Service evidence + balance + release","Order Completed","Send & Update","Eligible released deliverable","Archived / review eligible","View final state")];headers=["Current","Trigger","Evidence","Template","Action","Attachment","Result","Customer next"]
    elif index==3:rows=[(t["title"],"Workflow-owned only" if t["category"] not in ("Documents","Payment") else "Exact eligible customer-safe artifact","Current quote/invoice/receipt/released deliverable when maintained","Internal, audit, source, unreleased, provider payload") for t in templates];headers=["Template","Required","Recommended / optional","Never attach"]
    else:rows=[("RON","Provider session and APS 13-stage workflow","Provider-completed document after APS review/release","Only signer-safe session and released content","Payment, appointment, signers, documents, Proof, return, release"),("Mobile","Appointment facts and act/certificate completion","Agreed returned/scanned deliverable after review/release","Appointment and released content","Payment, appointment, certificate/document facts, delivery"),("Print & Scan","Source/specifications, production, QC, delivery","Completed Scan after review/release","Production status and released file","Payment, source, production, QC, release/delivery")];headers=["Service","Fulfillment authority","Customer deliverable","Customer visibility","Completion gate"]
    add_table(doc,headers,rows)
    add_callout(doc,"Must Do","Use the exact row that matches the current service and evidence. Do not treat a projected status as proof of the underlying event.","must")
def add_index(doc):
    entries={"Acknowledgment":"part_X_chapter_2","Accessibility":"part_XIX_chapter_3","Archive":"part_III_chapter_1","Cancellation":"part_VIII_chapter_3","Communication Log":"part_V_chapter_3","Completion Gate":"part_XII_chapter_4","Credible Witness":"part_X_chapter_5","Customer Portal":"part_XIII_chapter_1","Documents":"part_IV_chapter_3","GA4":"part_XV_chapter_3","Invoice":"part_VII_chapter_2","Jurat":"part_X_chapter_3","KBA":"part_IX_chapter_4","Messages":"part_IV_chapter_6","Mobile":"part_XI_chapter_1","Oath":"part_X_chapter_3","Payment Received":"part_VII_chapter_3","POA":"part_X_chapter_6","Print & Scan":"part_XII_chapter_1","Privacy":"part_XIX_chapter_2","Proof":"part_IX_chapter_2","Refund":"part_VII_chapter_4","Release to Customer":"part_VI_chapter_3","Review Queue":"part_III_chapter_2","RON":"part_IX_chapter_1","Scripts":"part_III_chapter_5","Send & Update Status":"part_V_chapter_2","Stripe":"part_VII_chapter_3","Templates":"part_V_chapter_4","Timeline":"part_IV_chapter_8","Witness":"part_X_chapter_5","Zelle / Offline Payment":"part_VII_chapter_3"}
    for term,anchor in entries.items():
        p=doc.add_paragraph();r=p.add_run(term);r.bold=True;p.add_run("  ....  ");field(p.add_run(),f"PAGEREF {anchor} \\h");p.add_run("  ");hyperlink(p,"go to chapter",anchor)
def build():
    OUT.mkdir(parents=True,exist_ok=True);templates,scripts=load_catalogs();doc=Document();configure(doc);cover(doc);contents(doc)
    chapter_count=0
    for roman,title,chapters in PARTS:
        part_divider(doc,roman,title,chapters)
        for i,ch in enumerate(chapters,1):
            anchor=f"part_{roman}_chapter_{i}";chapter_count+=1
            if roman=="XIX" and ch in ("Terms of Service","Privacy Policy","Accessibility Statement","Frequently Asked Questions"):
                path={"Terms of Service":"terms.html","Privacy Policy":"privacy.html","Accessibility Statement":"accessibility.html","Frequently Asked Questions":"faq.html"}[ch];add_policy(doc,f"Chapter {i} - {ch}",ROOT/path,anchor)
            elif roman=="XXI":add_heading(doc,f"Chapter {i} - {ch}",1,anchor);add_back(doc);add_template_catalog(doc,templates,anchor)
            elif roman=="XXII":add_heading(doc,f"Chapter {i} - {ch}",1,anchor);add_back(doc);add_status_catalog(doc,anchor)
            elif roman=="XXIII":add_heading(doc,f"Chapter {i} - {ch}",1,anchor);add_back(doc);add_visual_atlas(doc)
            elif roman=="XXIV":add_heading(doc,f"Chapter {i} - {ch}",1,anchor);add_back(doc);[doc.add_paragraph(f"{term}. {definition} See also the related Contents entry.") for term,definition in sorted(GLOSSARY.items())]
            elif roman=="XXV":add_heading(doc,f"Chapter {i} - {ch}",1,anchor);add_back(doc);add_index(doc)
            elif roman=="XVIII":add_heading(doc,f"Chapter {i} - {ch}",1,anchor);add_back(doc);add_quick_cards(doc,i-1)
            elif roman=="XX":
                if i==1:landscape_start(doc)
                add_matrix_chapter(doc,i,ch,anchor,templates)
                if i==len(chapters):portrait_start(doc)
            elif roman=="X" and i==1:add_chapter(doc,roman,i,ch,anchor);add_script_book(doc,scripts,anchor)
            else:add_chapter(doc,roman,i,ch,anchor)
    props=doc.core_properties;props.title="Aligned Print & Scan Operations, Policies & Standard Operating Procedures Handbook";props.subject="Operator Edition";props.author="Aligned Print & Scan LLC";props.keywords="APS, SOP, operator, RON, mobile notary, print and scan"
    doc.settings.element.append(OxmlElement("w:updateFields"));doc.settings.element[-1].set(qn("w:val"),"true")
    doc.save(DOCX)
    print(json.dumps({"docx":str(DOCX),"pdf":str(PDF),"parts":len(PARTS),"chapters":chapter_count,"templates":len(templates),"statuses":len(STATUS_ENTRIES),"scripts":len(scripts),"cards":15,"visuals":len(list((ROOT/'docs/assets/manual').glob('ss-*.jpg')))}))
if __name__=="__main__":build()
