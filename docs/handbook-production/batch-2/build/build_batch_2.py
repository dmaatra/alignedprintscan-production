#!/usr/bin/env python3
from pathlib import Path
import re
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
FIG = ROOT / "figures"
OUT.mkdir(parents=True, exist_ok=True)
FIG.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "APS_Operator_Handbook_Batch_2_Parts_V-VIII.docx"
SOURCE = ROOT / "PARTS_V-VIII.md"

NAVY = "161C4D"; GOLD = "C8A96B"; IVORY = "F6F3EE"; GREY = "6B6D78"
WHITE = "FFFFFF"; INK = "2D2D2D"; RED = "982D27"; GREEN = "2F6F58"

def font(run, name="Montserrat", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold; run.italic = italic

def shade(cell, color):
    pr = cell._tc.get_or_add_tcPr(); node = pr.find(qn("w:shd")) or OxmlElement("w:shd")
    node.set(qn("w:fill"), color)
    if node.getparent() is None: pr.append(node)

def margins(cell):
    pr = cell._tc.get_or_add_tcPr(); node = pr.first_child_found_in("w:tcMar")
    if node is None: node = OxmlElement("w:tcMar"); pr.append(node)
    for side, value in (("top", 100), ("start", 120), ("bottom", 100), ("end", 120)):
        item = node.find(qn("w:" + side)) or OxmlElement("w:" + side)
        item.set(qn("w:w"), str(value)); item.set(qn("w:type"), "dxa")
        if item.getparent() is None: node.append(item)

def bookmark(p, name, ident):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(ident)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(ident))
    p._p.insert(0, start); p._p.append(end)

def link(p, text, target):
    h = OxmlElement("w:hyperlink"); h.set(qn("w:anchor"), target); h.set(qn("w:history"), "1")
    r = OxmlElement("w:r"); rp = OxmlElement("w:rPr"); c = OxmlElement("w:color"); c.set(qn("w:val"), NAVY)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rp.extend((c, u)); r.append(rp)
    t = OxmlElement("w:t"); t.text = text; r.append(t); h.append(r); p._p.append(h)

def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.22
    for bit in re.split(r"(\*\*.*?\*\*)", text):
        if not bit: continue
        bold = bit.startswith("**") and bit.endswith("**")
        font(p.add_run(bit[2:-2] if bold else bit), bold=bold)
    return p

def add_table(doc, rows):
    table = doc.add_table(rows=0, cols=len(rows[0])); table.autofit = False
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, value in enumerate(row):
            cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cells[ci])
            shade(cells[ci], NAVY if ri == 0 else (WHITE if ri % 2 else IVORY))
            p = cells[ci].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            font(p.add_run(value), size=8.4, color=WHITE if ri == 0 else INK, bold=ri == 0)
        if ri == 0:
            pr = table.rows[0]._tr.get_or_add_trPr(); hdr = OxmlElement("w:tblHeader"); hdr.set(qn("w:val"), "true"); pr.append(hdr)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def draw_flow(path, title, nodes, decisions=()):
    w, h = 1500, 420; im = Image.new("RGB", (w, h), "#F6F3EE"); d = ImageDraw.Draw(im)
    try: title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34); body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 23)
    except OSError: title_font = body_font = ImageFont.load_default()
    d.text((50, 28), title, fill="#161C4D", font=title_font)
    gap = 24; usable = w - 100; bw = (usable - gap * (len(nodes) - 1)) / len(nodes); y = 150
    for i, label in enumerate(nodes):
        x = 50 + i * (bw + gap); box = (x, y, x + bw, y + 145)
        if i in decisions:
            cx = x + bw / 2; cy = y + 72; pts = [(cx, y), (x + bw, cy), (cx, y + 144), (x, cy)]
            d.polygon(pts, fill="#FFF8E6", outline="#C8A96B", width=5)
        else: d.rounded_rectangle(box, radius=18, fill="#FFFFFF", outline="#C8A96B", width=5)
        words = label.split(); lines = []; line = ""
        for word in words:
            trial = (line + " " + word).strip()
            if d.textlength(trial, font=body_font) > bw - 28 and line: lines.append(line); line = word
            else: line = trial
        if line: lines.append(line)
        ty = y + 48 - (len(lines) - 1) * 14
        for line in lines:
            tw = d.textlength(line, font=body_font); d.text((x + (bw - tw) / 2, ty), line, fill="#161C4D", font=body_font); ty += 30
        if i < len(nodes) - 1:
            ax = x + bw + 4; nx = x + bw + gap - 5; ay = y + 72
            d.line((ax, ay, nx, ay), fill="#161C4D", width=5); d.polygon([(nx, ay), (nx - 18, ay - 10), (nx - 18, ay + 10)], fill="#161C4D")
    im.save(path)

FIGURES = {
    18: ("fig-15-ron-workflow.png", "Remote Online Notary Workflow Map", ["Request review", "Quote / payment", "Appointment readiness", "Proof session", "Review / release", "Closeout"], (2,)),
    21: ("fig-16-ron-failure-decision.png", "RON Failure / Alternative Path", ["Failure identified", "Safe retry?", "Reuse transaction", "Alternative eligible?", "Reschedule / Mobile", "Record outcome"], (1, 3)),
    22: ("fig-17-mobile-workflow.png", "Mobile Notary Workflow Map", ["Request / location", "Route / quote", "Payment / appointment", "Travel", "At the table", "Delivery / closeout"], (2,)),
    26: ("fig-18-print-scan-workflow.png", "Print & Scan Workflow Map", ["File received", "Page count / specs", "Quote / payment", "Production", "QC", "Delivery / closeout"], (1,)),
}

for _, (name, title, nodes, decisions) in FIGURES.items(): draw_flow(FIG / name, title, nodes, decisions)

doc = Document(); sec = doc.sections[0]
sec.top_margin = Inches(.8); sec.bottom_margin = Inches(.75); sec.left_margin = Inches(.85); sec.right_margin = Inches(.85)
sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35); sec.different_first_page_header_footer = True
normal = doc.styles["Normal"]; normal.font.name = "Montserrat"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Montserrat"); normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.22
for name, size, before, after in (("Heading 1", 23, 18, 10), ("Heading 2", 17, 16, 8), ("Heading 3", 12.5, 12, 5)):
    s = doc.styles[name]; s.font.name = "Playfair Display" if name != "Heading 3" else "Montserrat"; s._element.rPr.rFonts.set(qn("w:ascii"), s.font.name); s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(NAVY); s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True
header = sec.header.paragraphs[0]; font(header.add_run("ALIGNED PRINT & SCAN  |  OPERATOR HANDBOOK · PARTS V–VIII"), size=8, color=GREY, bold=True)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(footer.add_run("Release 10 Certified Operations Edition  ·  Batch 2"), size=8, color=GREY)

for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("ALIGNED PRINT & SCAN"), size=13, color=GOLD, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("OPERATOR\nHANDBOOK"), name="Playfair Display", size=34, color=NAVY, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("Release 10 Certified Operations Edition"), size=14, color=GREY, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(55); font(p.add_run("BATCH 2  |  PARTS V–VIII  |  CHAPTERS 15–28"), size=10, color=GOLD, bold=True)
doc.add_page_break()

lines = SOURCE.read_text().splitlines(); headings = []
for line in lines:
    if line.startswith("# PART "): headings.append((1, line[2:]))
    elif line.startswith("## Chapter "): headings.append((2, line[3:]))
    elif line.startswith("### "): headings.append((3, line[4:]))
doc.add_heading("Contents", 0); add_para(doc, "Each entry links to the controlled Part, Chapter, or numbered subsection in this Batch 2 publication.")
for idx, (level, title) in enumerate(headings, 1):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches({1: 0, 2: .2, 3: .45}[level]); p.paragraph_format.space_after = Pt(2)
    link(p, title, f"toc_{idx}")
doc.add_page_break()

i = 0; bookmark_id = 1000; heading_index = 0; current_ch = None
while i < len(lines):
    line = lines[i].rstrip()
    if not line: i += 1; continue
    if line.startswith("# PART "):
        doc.add_page_break(); p = doc.add_heading(line[2:], 0); heading_index += 1; bookmark(p, f"toc_{heading_index}", bookmark_id); bookmark_id += 1
        i += 1; continue
    if line.startswith("## Chapter "):
        m = re.search(r"Chapter (\d+)", line); current_ch = int(m.group(1)); p = doc.add_heading(line[3:], 1); p.paragraph_format.page_break_before = True
        heading_index += 1; bookmark(p, f"toc_{heading_index}", bookmark_id); bookmark_id += 1
        add_para(doc, "Purpose: perform this workflow from verified APS evidence while preserving customer, legal, notarial, financial, and provider boundaries.")
        learn = add_para(doc, "YOU WILL LEARN TO")
        for run in learn.runs: font(run, size=9.5, color=NAVY, bold=True)
        add_para(doc, "Recognize readiness, perform the maintained operator action, communicate at the correct time, stop at blockers, and leave an auditable record.", style="List Bullet")
        if current_ch in FIGURES:
            fn, title, _, _ = FIGURES[current_ch]; pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER; pic.add_run().add_picture(str(FIG / fn), width=Inches(6.3))
            for drawing in pic._p.xpath(".//wp:docPr"): drawing.set("descr", title + " with connected workflow steps and decision gates")
            cap = doc.add_paragraph(title); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; font(cap.runs[0], size=8.5, color=GREY, italic=True)
        i += 1; continue
    if line.startswith("### "):
        p = doc.add_heading(line[4:], 2); heading_index += 1; bookmark(p, f"toc_{heading_index}", bookmark_id); bookmark_id += 1; i += 1; continue
    if line.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].startswith("|"):
            vals = [x.strip() for x in lines[i].strip("|").split("|")]
            if not all(set(x) <= set("-: ") for x in vals): rows.append(vals)
            i += 1
        if rows: add_table(doc, rows)
        continue
    if line.startswith("- "): add_para(doc, line[2:], style="List Bullet"); i += 1; continue
    if re.match(r"^\d+\. ", line): add_para(doc, re.sub(r"^\d+\. ", "", line), style="List Number"); i += 1; continue
    text = line; i += 1
    while i < len(lines) and lines[i] and not lines[i].startswith(("#", "|", "- ")) and not re.match(r"^\d+\. ", lines[i]): text += " " + lines[i].strip(); i += 1
    if text.startswith("**STOP"):
        p = add_para(doc, text); p.paragraph_format.left_indent = Inches(.15); p.paragraph_format.right_indent = Inches(.1)
        pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "FCE8E6"); pr.append(shd)
    else: add_para(doc, text)

doc.core_properties.title = "APS Operator Handbook — Batch 2 Parts V–VIII"
doc.core_properties.author = "Aligned Print & Scan"
settings = doc.settings._element; update = OxmlElement("w:updateFields"); update.set(qn("w:val"), "true"); settings.append(update)
doc.save(DOCX)
print({"docx": str(DOCX), "chapters": 14, "toc_entries": len(headings), "figures": len(FIGURES)})
