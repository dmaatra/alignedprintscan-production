#!/usr/bin/env python3
from pathlib import Path
import json, sys
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE=Path(__file__).resolve(); ROOT=HERE.parents[1]; PROD=ROOT.parent; sys.path.insert(0,str(PROD/"shared"))
from handbook_builder import *

OUT=ROOT/"output"; FIG=ROOT/"figures"; OUT.mkdir(parents=True,exist_ok=True); FIG.mkdir(parents=True,exist_ok=True)
DOCX=OUT/"APS_Operator_Handbook_Batch_3_Part_IX_Loan_Signing.docx"
SOURCE=ROOT/"PART_IX.md"; DOCUMENTS=json.loads((ROOT/"loan_documents.json").read_text())

FIGURES={
29:("fig-19-loan-signing-workflow.png","Loan Signing Workflow Map",["Assignment review","Package readiness","Print / appointment","Signing","Scanback / approval","Return / closeout"],(1,)),
35:("fig-20-loan-signing-ten-gates.png","Loan Signing Ten-Gate Pipeline",["1–2 Review / price","3–4 Package / print","5 Appointment","6 Signing","7 Scan / approval","8–10 Return / resolve / complete"],(2,4)),
36:("fig-21-package-version-decision.png","Package Version Decision Tree",["File received","Replacement explicit?","Mark active / preserve old","Instructions conflict?","Hold / contact orderer","Print / QC"],(1,3)),
37:("fig-22-signer-question-guide.png","Signer Question Decision Guide",["Signer asks","Procedural fact?","Point to printed text","Meaning / terms?","Refer to designated contact","Record outcome"],(1,3)),
38:("fig-23-scanback-approval-return.png","Scanback → Approval → Return Decision Tree",["Scanbacks required?","Scan / QC / upload","Approval required?","Hold securely","Approved","Return / tracking"],(0,2)),
39:("fig-24-loan-signing-exception.png","Loan Signing Exception Decision Tree",["Exception","Safe to continue?","Correct / complete","Stop / contact","Operational + financial resolution","Closeout"],(1,))
}
for _,(fn,title,nodes,decisions) in FIGURES.items(): flow(FIG/fn,title,nodes,decisions)

def catalog(doc,chapter):
    if chapter!=31:return
    doc.add_heading("31.1 Critical Notarization Rule",2)
    para(doc,"COMMONLY NOTARIZED DOES NOT MEAN ALWAYS NOTARIZE. Actual authority comes from the actual document, its actual certificate, applicable law, and authorized assignment instructions. Never add a certificate merely because a similar document is often notarized.")
    for i,item in enumerate(DOCUMENTS,2):
        heading=doc.add_heading(f"31.{i} {item['name']}",2)
        heading.paragraph_format.page_break_before=True
        table(doc,[["Reference field","Maintained training guidance"],["Common package(s)",item["packages"]],["What it is",item["what"]],["Why it appears",item["why"]],["Signer interaction",item["interaction"]],["Operator focus",item["focus"]],["Commonly notarized?",item["notarized"]],["Presentation guidance",item["guidance"]],["Do not",item["do_not"]],["If signer has a question",item["questions"]],["Source",item["source"]]])

doc=new_document("BATCH 3 · PART IX")
cover(doc,"Batch 3","PART IX  |  CHAPTERS 29–40")
lines=SOURCE.read_text().splitlines(); items=headings(lines)
add_contents(doc,items)
figure_map={ch:(FIG/fn,title) for ch,(fn,title,_,_) in FIGURES.items()}
parse_markdown(doc,lines,figure_map,catalog)
doc.add_page_break(); doc.add_heading("Loan Document Source Register",1)
for line in (ROOT/"LOAN_DOCUMENT_SOURCE_REGISTER.md").read_text().splitlines():
    if line.startswith("## "): doc.add_heading(line[3:],2)
    elif line.startswith("- "): para(doc,line[2:],style="List Bullet")
    elif line and not line.startswith("#"): para(doc,line)
doc.core_properties.title="APS Operator Handbook — Batch 3 Part IX Loan Signing"; doc.core_properties.author="Aligned Print & Scan"
u=OxmlElement("w:updateFields"); u.set(qn("w:val"),"true"); doc.settings._element.append(u); doc.save(DOCX)
print({"docx":str(DOCX),"chapters":12,"document_entries":len(DOCUMENTS),"toc_entries":len(items),"figures":len(FIGURES)})
