#!/usr/bin/env python3
from pathlib import Path
import importlib.util, json, re, sys
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE=Path(__file__).resolve(); ROOT=HERE.parents[1]; PROD=ROOT.parent; REPO=PROD.parents[1]
sys.path.insert(0,str(PROD/"shared"))
from handbook_builder import *

OUT=ROOT/"output"; OUT.mkdir(parents=True,exist_ok=True)
DOCX=OUT/"APS_Operator_Handbook_Release_10_Certified_Operations_Edition.docx"

sources=[
 PROD/"batch-1/part-01/PART_I.md", PROD/"batch-1/part-02/PART_II.md", PROD/"batch-1/part-03/PART_III.md", PROD/"batch-1/part-04/PART_IV.md",
 PROD/"batch-2/PARTS_V-VIII.md", PROD/"batch-3/PART_IX.md", PROD/"batch-4/PARTS_X_XI.md", PROD/"batch-5/PARTS_XII_XIII.md",
]
lines=[]
for src in sources:
    part=src.read_text().replace("# Part ","# PART ").splitlines()
    lines.extend(part); lines.append("")

def load(path,name):
    spec=importlib.util.spec_from_file_location(name,path); mod=importlib.util.module_from_spec(spec); spec.loader.exec_module(mod); return mod
b3=load(PROD/"batch-3/build/build_batch_3.py","handbook_batch3")
b4=load(PROD/"batch-4/build/build_batch_4.py","handbook_batch4")

figure_candidates={}
for batch in ("batch-1","batch-2","batch-3","batch-4","batch-5"):
    for path in sorted((PROD/batch/"figures").glob("*.png")):
        match=re.match(r"fig-(\d+)-",path.name)
        if match: figure_candidates[int(match.group(1))]=path
chapter_figure={1:1,3:2,4:3,7:4,8:5,9:7,10:8,11:10,12:9,14:14,18:15,21:16,22:17,26:18,29:19,35:20,36:21,37:22,38:23,39:24,41:25,43:26,45:27,48:28,49:29,50:30}
titles={p.stem:p.stem.replace("fig-","").replace("-"," ").title() for p in figure_candidates.values()}
figure_map={ch:(figure_candidates[num],titles[figure_candidates[num].stem]) for ch,num in chapter_figure.items() if num in figure_candidates}

ss_by_ch={4:["ss-100-portal-overview.jpg","ss-101-portal-documents.jpg"],5:["ss-001-requests.jpg","ss-010-review-queue.jpg"],6:["ss-020-request-overview.jpg"],9:["ss-030-document-lifecycle.jpg"],10:["ss-031-private-output-review.jpg"],12:["ss-032-release-to-customer.jpg"]}
extra_figs={1:[3],2:[5,6],14:[11,12,13]}

def hook(doc,chapter):
    if chapter==31: b3.catalog(doc,chapter)
    if chapter=="TEMPLATE_DIRECTORY": b4.add_templates(doc)
    if chapter=="SCRIPT_DIRECTORY": b4.add_scripts(doc)
    if isinstance(chapter,int):
        for num in extra_figs.get(chapter,[]):
            path=figure_candidates.get(num)
            if path: add_figure(doc,path,titles[path.stem])
        for name in ss_by_ch.get(chapter,[]):
            path=REPO/"docs/assets/manual"/name
            if path.exists(): add_figure(doc,path,"Governed synthetic APS interface — "+name.replace(".jpg","").replace("-"," "))

doc=new_document("COMPLETE · PARTS I–XIII")
cover(doc,"Certified Operations Edition","PARTS I–XIII  |  CHAPTERS 1–53")
doc.add_heading("Document Control",0)
para(doc,"Certified baseline: APS Release 10. Edition: Certified Operations Edition. Publication authority: OWNER APPROVED — LOCKED specification, PR #107. This editable master is assembled from controlled chapter source, not from unrelated application code or live production data.")
doc.add_heading("How to Use This Handbook",1)
para(doc,"LEARN in Parts I–IV; OPERATE in Parts V–XII; use Part XIII for QUICK REFERENCE. The detailed governing chapter controls whenever a Quick Reference is abbreviated. Stop and source-verify any material drift.")
doc.add_heading("Scope and Professional Boundaries",1)
para(doc,"Operators coordinate verified APS work within assigned authority. They do not give legal or financial advice, choose legal documents or notarial acts for customers, interpret loan terms, bypass fulfillment gates, expose protected information, or manufacture provider, payment, delivery, or audit outcomes.")
doc.add_heading("How to Read Workflow Diagrams",1)
para(doc,"Rectangles show controlled stages or actions; diamonds show decisions; arrows show permitted sequence. A branch is not permission to bypass the governing conditions. Text explanations remain the accessible authority.")
doc.add_page_break()
items=headings(lines); add_contents(doc,items)
parse_markdown(doc,lines,figure_map,hook)

doc.add_page_break(); doc.add_heading("List of Figures",0)
for num,path in sorted(figure_candidates.items()): para(doc,f"Figure {num:02d} — {titles[path.stem]}",style="List Bullet")
for title in ("List of Tables","List of Procedures","List of Checklists","List of Decision Trees"):
    doc.add_heading(title,0); para(doc,"Controlled entries are indexed through the numbered chapter/subsection headings and captions in this edition. Use the expanded Contents and PDF bookmark tree for direct navigation.")
doc.add_heading("Source Authority and Traceability",0)
para(doc,"Certified Release 10 behavior and current maintained repository authority control. Batch source registers preserve Texas notary, provider, lender/title/agency, template, script, financial, and loan-document qualification. SOURCE VERIFICATION REQUIRED items are never promoted to established facts.")
doc.add_heading("Revision History and Owner Approval Record",0)
para(doc,"Specification PR #107; Batch 1 PRs #108–109; Batch 2 PR #114; Batch 3 PR #115; Batch 4 PR #116; Batch 5 PR #117. Final assembly is published only after structural, accessibility, visual, navigation, inventory, source, and safety certification.")
doc.core_properties.title="APS Operator Handbook — Release 10 Certified Operations Edition"; doc.core_properties.author="Aligned Print & Scan"
u=OxmlElement("w:updateFields"); u.set(qn("w:val"),"true"); doc.settings._element.append(u); doc.save(DOCX)
print({"docx":str(DOCX),"parts":13,"chapters":53,"toc_entries":len(items),"figures":len(figure_candidates),"templates":len(b4.TEMPLATES),"scripts":len(b4.SCRIPTS)})
