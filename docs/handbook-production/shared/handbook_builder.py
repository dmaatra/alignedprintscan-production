from pathlib import Path
import re
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY="161C4D"; GOLD="C8A96B"; IVORY="F6F3EE"; GREY="6B6D78"; WHITE="FFFFFF"; INK="2D2D2D"

def font(run,name="Montserrat",size=10.5,color=INK,bold=False,italic=False):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),name); run._element.rPr.rFonts.set(qn("w:hAnsi"),name)
    run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic

def bookmark(p,name,ident):
    a=OxmlElement("w:bookmarkStart"); a.set(qn("w:id"),str(ident)); a.set(qn("w:name"),name); b=OxmlElement("w:bookmarkEnd"); b.set(qn("w:id"),str(ident)); p._p.insert(0,a); p._p.append(b)

def link(p,text,target):
    h=OxmlElement("w:hyperlink"); h.set(qn("w:anchor"),target); h.set(qn("w:history"),"1"); r=OxmlElement("w:r"); rp=OxmlElement("w:rPr"); c=OxmlElement("w:color"); c.set(qn("w:val"),NAVY); u=OxmlElement("w:u"); u.set(qn("w:val"),"single"); rp.extend((c,u)); r.append(rp); t=OxmlElement("w:t"); t.text=text; r.append(t); h.append(r); p._p.append(h)

def para(doc,text,style=None):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.22
    for bit in re.split(r"(\*\*.*?\*\*)",text):
        if bit: font(p.add_run(bit[2:-2] if bit.startswith("**") and bit.endswith("**") else bit),bold=bit.startswith("**") and bit.endswith("**"))
    return p

def table(doc,rows):
    t=doc.add_table(rows=0,cols=len(rows[0])); t.autofit=False
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        row_pr=t.rows[-1]._tr.get_or_add_trPr(); no_split=OxmlElement("w:cantSplit"); row_pr.append(no_split)
        for ci,value in enumerate(row):
            c=cells[ci]; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            pr=c._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),NAVY if ri==0 else (WHITE if ri%2 else IVORY)); pr.append(sh)
            mar=OxmlElement("w:tcMar")
            for side,val in (("top",100),("start",120),("bottom",100),("end",120)):
                n=OxmlElement("w:"+side); n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa"); mar.append(n)
            pr.append(mar); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(str(value)),size=8.2,color=WHITE if ri==0 else INK,bold=ri==0)
        if ri==0:
            pr=t.rows[0]._tr.get_or_add_trPr(); h=OxmlElement("w:tblHeader"); h.set(qn("w:val"),"true"); pr.append(h)
    doc.add_paragraph()

def flow(path,title,nodes,decisions=()):
    w,h=1500,420; im=Image.new("RGB",(w,h),"#F6F3EE"); d=ImageDraw.Draw(im)
    try: tf=ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf",34); bf=ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf",22)
    except OSError: tf=bf=ImageFont.load_default()
    d.text((50,28),title,fill="#161C4D",font=tf); gap=22; bw=(w-100-gap*(len(nodes)-1))/len(nodes); y=150
    for i,label in enumerate(nodes):
        x=50+i*(bw+gap); cx=x+bw/2; cy=y+72
        if i in decisions: d.polygon([(cx,y),(x+bw,cy),(cx,y+144),(x,cy)],fill="#FFF8E6",outline="#C8A96B",width=5)
        else: d.rounded_rectangle((x,y,x+bw,y+145),radius=18,fill="white",outline="#C8A96B",width=5)
        words=label.split(); lines=[]; line=""
        for word in words:
            trial=(line+" "+word).strip()
            if line and d.textlength(trial,font=bf)>bw-26: lines.append(line); line=word
            else: line=trial
        lines.append(line); ty=y+49-(len(lines)-1)*13
        for ln in lines: tw=d.textlength(ln,font=bf); d.text((x+(bw-tw)/2,ty),ln,fill="#161C4D",font=bf); ty+=28
        if i<len(nodes)-1:
            ax=x+bw+3; nx=x+bw+gap-4; d.line((ax,cy,nx,cy),fill="#161C4D",width=5); d.polygon([(nx,cy),(nx-16,cy-9),(nx-16,cy+9)],fill="#161C4D")
    im.save(path)

def new_document(batch_label):
    d=Document(); s=d.sections[0]; s.top_margin=Inches(.8); s.bottom_margin=Inches(.75); s.left_margin=Inches(.85); s.right_margin=Inches(.85); s.header_distance=Inches(.35); s.footer_distance=Inches(.35); s.different_first_page_header_footer=True
    n=d.styles["Normal"]; n.font.name="Montserrat"; n._element.rPr.rFonts.set(qn("w:ascii"),"Montserrat"); n.font.size=Pt(10.5); n.font.color.rgb=RGBColor.from_string(INK); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.22
    for name,size,before,after in (("Heading 1",23,18,10),("Heading 2",17,16,8),("Heading 3",12.5,12,5)):
        st=d.styles[name]; st.font.name="Playfair Display" if name!="Heading 3" else "Montserrat"; st._element.rPr.rFonts.set(qn("w:ascii"),st.font.name); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(NAVY); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    font(s.header.paragraphs[0].add_run("ALIGNED PRINT & SCAN  |  OPERATOR HANDBOOK · "+batch_label),size=8,color=GREY,bold=True)
    s.footer.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; font(s.footer.paragraphs[0].add_run("Release 10 Certified Operations Edition  ·  "+batch_label),size=8,color=GREY)
    return d

def cover(doc,batch_label,scope):
    for _ in range(6): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("ALIGNED PRINT & SCAN"),size=13,color=GOLD,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("OPERATOR\nHANDBOOK"),name="Playfair Display",size=34,color=NAVY,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("Release 10 Certified Operations Edition"),size=14,color=GREY,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55); font(p.add_run(batch_label.upper()+"  |  "+scope),size=10,color=GOLD,bold=True); doc.add_page_break()

def headings(lines):
    result=[]
    for line in lines:
        if line.startswith("# PART "): result.append((1,line[2:]))
        elif line.startswith("## Chapter "): result.append((2,line[3:]))
        elif line.startswith("### "): result.append((3,line[4:]))
    return result

def add_contents(doc,items):
    doc.add_heading("Contents",0); para(doc,"Part, Chapter, and numbered subsection entries link to the controlled publication headings.")
    for i,(level,title) in enumerate(items,1):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches({1:0,2:.2,3:.45}[level]); p.paragraph_format.space_after=Pt(2); link(p,title,f"toc_{i}")
    doc.add_page_break()

def add_figure(doc,path,title):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path),width=Inches(6.3))
    for drawing in p._p.xpath(".//wp:docPr"): drawing.set("descr",title+" with connected steps and decision gates")
    c=doc.add_paragraph(title); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(c.runs[0],size=8.5,color=GREY,italic=True)

def parse_markdown(doc,lines,figure_map=None,chapter_hook=None):
    i=0; bid=1000; hi=0; current=None; figure_map=figure_map or {}
    while i<len(lines):
        line=lines[i].rstrip()
        if not line: i+=1; continue
        if line in ("{{TEMPLATE_DIRECTORY}}","{{SCRIPT_DIRECTORY}}"):
            if chapter_hook: chapter_hook(doc,line.strip("{}"))
            i+=1; continue
        if line.startswith("# PART "):
            p=doc.add_heading(line[2:],0); hi+=1; bookmark(p,f"toc_{hi}",bid); bid+=1; i+=1; continue
        if line.startswith("## Chapter "):
            current=int(re.search(r"Chapter (\d+)",line).group(1)); p=doc.add_heading(line[3:],1); p.paragraph_format.page_break_before=True; hi+=1; bookmark(p,f"toc_{hi}",bid); bid+=1
            para(doc,"Purpose: perform this workflow from verified APS and assignment evidence while preserving legal, notarial, financial, privacy, and ordering-party boundaries.")
            learn=para(doc,"YOU WILL LEARN TO"); [font(r,size=9.5,color=NAVY,bold=True) for r in learn.runs]; para(doc,"Review the authoritative assignment, act within role, stop at blockers, communicate accurately, and leave a complete audit trail.",style="List Bullet")
            if current in figure_map: add_figure(doc,*figure_map[current])
            if chapter_hook: chapter_hook(doc,current)
            i+=1; continue
        if line.startswith("### "):
            p=doc.add_heading(line[4:],2); hi+=1; bookmark(p,f"toc_{hi}",bid); bid+=1; i+=1; continue
        if line.startswith("|"):
            rows=[]
            while i<len(lines) and lines[i].startswith("|"):
                vals=[x.strip() for x in lines[i].strip("|").split("|")]
                if not all(set(x)<=set("-: ") for x in vals): rows.append(vals)
                i+=1
            if rows: table(doc,rows)
            continue
        if line.startswith("- "): para(doc,line[2:],style="List Bullet"); i+=1; continue
        if re.match(r"^\d+\. ",line):
            p=para(doc,line); p.paragraph_format.left_indent=Inches(.22); p.paragraph_format.first_line_indent=Inches(-.22); i+=1; continue
        text=line; i+=1
        while i<len(lines) and lines[i] and not lines[i].startswith(("#","|","- ")) and not re.match(r"^\d+\. ",lines[i]): text+=" "+lines[i].strip(); i+=1
        para(doc,text)
    return hi
