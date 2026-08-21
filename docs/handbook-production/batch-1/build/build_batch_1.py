from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

ROOT=Path(__file__).resolve().parents[1]; REPO=ROOT.parents[2]
OUT=ROOT/'output'; OUT.mkdir(parents=True,exist_ok=True)
DOCX=OUT/'APS_Operator_Handbook_Batch_1_Parts_I-IV.docx'
NAVY='161C4D'; GOLD='C8A96B'; LIGHT='E8D28D'; IVORY='F6F3EE'; CREAM='F3EEE4'; GREY='6B6D78'; BORDER='E4E0D8'; RED='B33A3A'; GREEN='2F8F5B'; WHITE='FFFFFF'; INK='2D2D2D'
SOURCES=[ROOT/'front-matter/FRONT_MATTER.md',ROOT/'part-01/PART_I.md',ROOT/'part-02/PART_II.md',ROOT/'part-03/PART_III.md',ROOT/'part-04/PART_IV.md',ROOT/'quick-reference/BATCH_1_QUICK_REFERENCES.md',ROOT/'sources/BATCH_1_SOURCE_REGISTER.md']
FIG_BY_CH={1:['fig-01-aps-system-map','fig-03-portal-system-relationship'],2:['fig-05-customer-signer-orderer','fig-06-participant-role-relationship'],3:['fig-02-universal-request-lifecycle'],4:['fig-03-portal-system-relationship'],7:['fig-04-state-mental-map'],8:['fig-05-customer-signer-orderer','fig-06-participant-role-relationship'],9:['fig-07-document-lifecycle'],10:['fig-08-document-security-boundary'],11:['fig-10-ron-proof-document'],12:['fig-09-document-release-decision'],14:['fig-11-mobile-original-scan','fig-12-print-source-output','fig-13-loan-package-return','fig-14-service-document-comparison']}
SS_BY_CH={4:['ss-100-portal-overview.jpg','ss-101-portal-documents.jpg'],5:['ss-001-requests.jpg','ss-010-review-queue.jpg'],6:['ss-020-request-overview.jpg'],9:['ss-030-document-lifecycle.jpg'],10:['ss-031-private-output-review.jpg'],12:['ss-032-release-to-customer.jpg']}

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd')) or OxmlElement('w:shd'); shd.set(qn('w:fill'),color)
    if shd.getparent() is None: tcPr.append(shd)
def set_cell_margins(cell,top=100,start=120,bottom=100,end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); mar=tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); tcPr.append(mar)
    for k,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=mar.find(qn('w:'+k)) or OxmlElement('w:'+k); node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')
        if node.getparent() is None: mar.append(node)
def font(run,name='Montserrat',size=10.5,color=INK,bold=False,italic=False):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name); run._element.rPr.rFonts.set(qn('w:hAnsi'),name); run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic
def field(paragraph,instr):
    r=paragraph.add_run(); begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin'); ins=OxmlElement('w:instrText'); ins.set(qn('xml:space'),'preserve'); ins.text=instr; sep=OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate'); txt=OxmlElement('w:t'); txt.text='1'; end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end'); r._r.extend([begin,ins,sep,txt,end]); return r
def bookmark(paragraph,name,bid):
    start=OxmlElement('w:bookmarkStart'); start.set(qn('w:id'),str(bid)); start.set(qn('w:name'),name); end=OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'),str(bid)); paragraph._p.insert(0,start); paragraph._p.append(end)
def internal_link(paragraph,text,target):
    link=OxmlElement('w:hyperlink'); link.set(qn('w:anchor'),target); link.set(qn('w:history'),'1'); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); color=OxmlElement('w:color'); color.set(qn('w:val'),NAVY); underline=OxmlElement('w:u'); underline.set(qn('w:val'),'single'); rPr.extend([color,underline]); t=OxmlElement('w:t'); t.text=text; r.extend([rPr,t]); link.append(r); paragraph._p.append(link)
def new_numbering():
    root=doc.part.numbering_part.element; ids=[int(n.get(qn('w:numId'))) for n in root.findall(qn('w:num'))]; num_id=max(ids or [0])+1; num=OxmlElement('w:num'); num.set(qn('w:numId'),str(num_id)); aid=OxmlElement('w:abstractNumId'); aid.set(qn('w:val'),'0'); num.append(aid); override=OxmlElement('w:lvlOverride'); override.set(qn('w:ilvl'),'0'); start=OxmlElement('w:startOverride'); start.set(qn('w:val'),'1'); override.append(start); num.append(override); root.append(num); return num_id
def add_numbered(doc,text,num_id):
    p=add_para(doc,text); pPr=p._p.get_or_add_pPr(); numPr=OxmlElement('w:numPr'); ilvl=OxmlElement('w:ilvl'); ilvl.set(qn('w:val'),'0'); nid=OxmlElement('w:numId'); nid.set(qn('w:val'),str(num_id)); numPr.extend([ilvl,nid]); pPr.insert(0,numPr); return p
def page_break(doc): doc.add_page_break()
def add_para(doc,text,style=None):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.22
    # simple bold spans
    bits=re.split(r'(\*\*.*?\*\*)',text)
    for bit in bits:
        if not bit: continue
        bold=bit.startswith('**') and bit.endswith('**'); value=bit[2:-2] if bold else bit; font(p.add_run(value),bold=bold)
    return p
def add_callout(doc,label,text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.right_indent=Inches(.08); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(8)
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'FFF8E6' if label not in ('STOP',) else 'FCE9E7'); pPr.append(shd)
    border=OxmlElement('w:pBdr'); left=OxmlElement('w:left'); left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'22'); left.set(qn('w:color'),RED if label=='STOP' else GOLD); border.append(left); pPr.append(border)
    font(p.add_run(label+' — '),bold=True,color=RED if label=='STOP' else NAVY); font(p.add_run(text))
def add_table(doc,rows):
    cols=len(rows[0]); widths=[Inches(6.5/cols)]*cols; t=doc.add_table(rows=0,cols=cols); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].width=widths[i]; cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cells[i]); shade(cells[i],NAVY if ri==0 else (WHITE if ri%2 else IVORY)); p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(val),size=8.5,color=WHITE if ri==0 else INK,bold=ri==0)
        if ri==0:
            trPr=t.rows[0]._tr.get_or_add_trPr(); hdr=OxmlElement('w:tblHeader'); hdr.set(qn('w:val'),'true'); trPr.append(hdr)
    doc.add_paragraph().paragraph_format.space_after=Pt(2); return t
def add_figure(doc,path,caption):
    with Image.open(path) as im: portrait=im.height>im.width*1.15
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True; p.add_run().add_picture(str(path),width=Inches(4.55 if portrait else 6.3))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(10); font(c.add_run(caption),size=8.5,color=GREY,italic=True)

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35); sec.different_first_page_header_footer=True
styles=doc.styles
normal=styles['Normal']; normal.font.name='Montserrat'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Montserrat'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.22
for name,size,before,after in [('Heading 1',23,18,10),('Heading 2',17,16,8),('Heading 3',12.5,12,5)]:
    s=styles[name]; s.font.name='Playfair Display' if name!='Heading 3' else 'Montserrat'; s._element.rPr.rFonts.set(qn('w:ascii'),s.font.name); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(NAVY); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
# running header/footer
h=sec.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.LEFT; font(h.add_run('ALIGNED PRINT & SCAN  |  OPERATOR HANDBOOK · PARTS I–IV'),size=8,color=GREY,bold=True)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(f.add_run('Operator Handbook  ·  Release 10 Certified Operations Edition  ·  '),size=8,color=GREY); field(f,'PAGE')
# cover editorial pattern
for _ in range(6): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('ALIGNED PRINT & SCAN'),name='Montserrat',size=13,color=GOLD,bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('OPERATOR\nHANDBOOK'),name='Playfair Display',size=34,color=NAVY,bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run('Release 10 Certified Operations Edition'),size=14,color=GREY,bold=True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(24); font(p.add_run('Remote Online Notary  ·  Mobile Notary\nPrint & Scan  ·  Loan Signing'),size=11,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55); font(p.add_run('BATCH 1 PREVIEW  |  PARTS I–IV  |  CHAPTERS 1–14'),size=10,color=GOLD,bold=True)
page_break(doc)
# TOC framework
doc.add_heading('Contents',0); add_para(doc,'This visible Batch 1 contents list links to the true chapter headings. The Word master also preserves the heading hierarchy used for PDF bookmarks and later full-edition assembly.')
TOC=[('Part I — Aligned Print & Scan',None),('Chapter 1 — APS Mission, Services & Operating Model','ch01'),('Chapter 2 — Roles, Authority & Professional Boundaries','ch02'),('Part II — Customer & Portal Operations',None),('Chapter 3 — The Customer Journey','ch03'),('Chapter 4 — Customer Portal Operations','ch04'),('Part III — Admin Operations',None),('Chapter 5 — Admin Navigation & Daily Triage','ch05'),('Chapter 6 — The Request Workspace','ch06'),('Chapter 7 — Understanding Status, Next Action & Gates','ch07'),('Part IV — Customers, Participants, Documents & Records',None),('Chapter 8 — Customers, Signers, Witnesses & Participants','ch08'),('Chapter 9 — Document Mental Model, Customer Uploads & Review','ch09'),('Chapter 10 — Document Classification & Safe Handling','ch10'),('Chapter 11 — Completed & Proof-Produced Documents','ch11'),('Chapter 12 — Document Release, Delivery & Audit','ch12'),('Chapter 13 — Document Replacement, Mistakes & Recovery','ch13'),('Chapter 14 — Service-Specific Document Lifecycles','ch14')]
for label,target in TOC:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.left_indent=Inches(.2 if target else 0); font(p.add_run(label) if not target else p.add_run(''),size=9.5,bold=not target,color=NAVY)
    if target: internal_link(p,label,target)
page_break(doc)
doc.add_heading('Lists and navigation',0)
for title,items in [('Figures',[f'Figure {i:02d} — {v[0]}' for i,v in enumerate(__import__('generate_figures').FIGURES.values(),1)]),('Tables',['Service relationship','Role comparison','Customer stages','Portal troubleshooting','Admin navigation','Request workspace tabs','Participant roles','Document classifications','Recovery and communication sequence']),('Procedures',[f'Procedure {x}' for x in ['1.1 Orient a new request','2.1 Decide whether you may act','3.1 Identify customer stage','4.1 Investigate portal question','5.1 Daily opening triage','5.2 Daily closeout','6.1 Work a Next Action','7.1 Evaluate a gate','8.1 Add/correct participant','9.1 Review upload','10.1 Handle safely','11.1 Missing Proof output','12.1 Release document','13.1 Wrong release','14.1 Identify document lane']]),('Checklists',['Completed-document review','Document Review','Document Release','Document Mistake Recovery'])]:
    doc.add_heading(title,1)
    for item in items: add_para(doc,item,style='List Bullet')
page_break(doc)

def parse_file(path):
    lines=path.read_text().splitlines(); i=0; current_ch=None; num_id=None; previous_numbered=False
    while i<len(lines):
        line=lines[i].rstrip();
        if not line: i+=1; continue
        if line.startswith('|'):
            rows=[]
            while i<len(lines) and lines[i].startswith('|'):
                vals=[x.strip() for x in lines[i].strip('|').split('|')]
                if not all(set(v)<=set('-: ') for v in vals): rows.append(vals)
                i+=1
            if rows: add_table(doc,rows)
            continue
        if line.startswith('# '):
            p=doc.add_heading(line[2:],0); p.paragraph_format.page_break_before=('Front Matter' not in line); i+=1; previous_numbered=False; continue
        if line.startswith('## '):
            title=line[3:]; m=re.match(r'Chapter (\d+)',title); current_ch=int(m.group(1)) if m else current_ch
            p=doc.add_heading(title,1); p.paragraph_format.page_break_before=bool(m)
            if m: bookmark(p,f'ch{current_ch:02d}',100+current_ch)
            if m:
                for slug in FIG_BY_CH.get(current_ch,[]):
                    fp=ROOT/'figures'/f'{slug}.png'
                    if fp.exists(): add_figure(doc,fp,'Figure '+slug.split('-')[1]+' — '+slug.replace('fig-'+slug.split('-')[1]+'-','').replace('-',' ').title())
                for ss in SS_BY_CH.get(current_ch,[]):
                    fp=REPO/'docs/assets/manual'/ss
                    if fp.exists(): add_figure(doc,fp,'APS governed synthetic screenshot — '+ss.replace('.jpg','').replace('-',' '))
            i+=1; previous_numbered=False; continue
        if line.startswith('### '): doc.add_heading(line[4:],2); i+=1; previous_numbered=False; continue
        if re.match(r'^\d+\. ',line):
            if not previous_numbered: num_id=new_numbering()
            add_numbered(doc,re.sub(r'^\d+\. ','',line),num_id); i+=1; previous_numbered=True; continue
        if line.startswith('- '): add_para(doc,line[2:],style='List Bullet'); i+=1; previous_numbered=False; continue
        m=re.match(r'^\*\*(OPERATOR RULE|STOP|DECISION|CUSTOMER COMMUNICATION|QUICK TIP)(?: — |:)(.*)\**$',line)
        if m: add_callout(doc,m.group(1),m.group(2).strip('* ')); i+=1; previous_numbered=False; continue
        # combine adjacent prose lines
        text=line; i+=1
        while i<len(lines) and lines[i] and not lines[i].startswith(('#','|','- ')) and not re.match(r'^\d+\. ',lines[i]): text+=' '+lines[i].strip(); i+=1
        previous_numbered=False
        if text.startswith('**STOP'):
            cleaned=text.replace('**',''); label,_,rest=cleaned.partition('—'); add_callout(doc,'STOP',rest.strip() or label.replace('STOP','').strip(' :'))
        elif text.startswith('**OPERATOR RULE'):
            cleaned=text.replace('**',''); label,_,rest=cleaned.partition('—'); add_callout(doc,'OPERATOR RULE',rest.strip())
        else: add_para(doc,text)

for src in SOURCES: parse_file(src)
settings=doc.settings._element; upd=OxmlElement('w:updateFields'); upd.set(qn('w:val'),'true'); settings.append(upd)
doc.core_properties.title='Aligned Print & Scan Operator Handbook — Batch 1 Parts I–IV'; doc.core_properties.subject='Release 10 Certified Operations Edition'; doc.core_properties.author='Aligned Print & Scan'; doc.core_properties.comments='Owner/editorial review preview; Chapters 1–14 only.'
doc.save(DOCX)
print({'docx':str(DOCX),'sources':len(SOURCES),'chapters':14})
